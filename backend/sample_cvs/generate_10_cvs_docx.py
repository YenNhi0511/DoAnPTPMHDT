"""
Script tạo 10 CV mẫu dạng DOCX cho các lĩnh vực khác nhau
Chạy: python generate_10_cvs_docx.py
Cần cài: pip install python-docx faker
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from faker import Faker

fake = Faker('vi_VN')

def sanitize_filename(name):
    """Chuyển tên thành tên file hợp lệ"""
    import re
    # Loại bỏ dấu tiếng Việt và ký tự đặc biệt
    name = name.replace(' ', '_')
    name = re.sub(r'[^\w\-_\.]', '', name)
    return name

def create_cv_docx(filepath, name, email, phone, address, experience_years, skills, position, group, profession):
    """Tạo CV dạng DOCX"""
    doc = Document()
    
    # Thiết lập style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    
    # ========== HEADER ==========
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run(name.upper())
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 139)  # Dark blue
    
    # Thông tin liên hệ
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.add_run(f"📧 {email} | 📱 {phone} | 📍 {address}")
    contact.runs[0].font.size = Pt(10)
    
    doc.add_paragraph()  # Khoảng trống
    
    # ========== MỤC TIÊU NGHỀ NGHIỆP ==========
    doc.add_paragraph().add_run("MỤC TIÊU NGHỀ NGHIỆP").bold = True
    doc.add_paragraph().add_run(f"Tìm kiếm cơ hội phát triển trong lĩnh vực {profession} với vị trí {position}. Mong muốn đóng góp kinh nghiệm {experience_years} năm và kỹ năng chuyên môn để phát triển sự nghiệp trong môi trường chuyên nghiệp.")
    doc.add_paragraph()  # Khoảng trống
    
    # ========== KINH NGHIỆM LÀM VIỆC ==========
    doc.add_paragraph().add_run("KINH NGHIỆM LÀM VIỆC").bold = True
    
    # Tạo 2-3 công việc trước đó
    num_jobs = fake.random_int(2, 3)
    for i in range(num_jobs):
        years_ago = experience_years - (num_jobs - i) + 1
        if years_ago < 1:
            years_ago = 1
        
        # Tên công ty và vị trí phù hợp với lĩnh vực
        company_name = fake.company()
        job_title = position if i == 0 else f"{position} (Trước đây)"
        
        p = doc.add_paragraph()
        p.add_run(f"{job_title}").bold = True
        p.add_run(f" | {company_name} | {years_ago} năm")
        
        # Mô tả công việc
        doc.add_paragraph(f"• {fake.sentence()}", style='List Bullet')
        doc.add_paragraph(f"• {fake.sentence()}", style='List Bullet')
        doc.add_paragraph(f"• {fake.sentence()}", style='List Bullet')
        
        if i < num_jobs - 1:
            doc.add_paragraph()  # Khoảng trống giữa các công việc
    
    doc.add_paragraph()  # Khoảng trống
    
    # ========== HỌC VẤN ==========
    doc.add_paragraph().add_run("HỌC VẤN").bold = True
    
    # Tạo học vấn phù hợp với lĩnh vực
    if 'Công nghệ Thông tin' in group or 'IT' in group:
        major = fake.random_element(elements=('Công nghệ Thông tin', 'Khoa học Máy tính', 'Kỹ thuật Phần mềm'))
    elif 'Marketing' in group or 'PR' in group:
        major = fake.random_element(elements=('Marketing', 'Quan hệ Công chúng', 'Truyền thông'))
    elif 'Kế toán' in group or 'Tài chính' in group:
        major = fake.random_element(elements=('Kế toán', 'Tài chính Ngân hàng', 'Quản trị Kinh doanh'))
    elif 'Nhân sự' in group:
        major = fake.random_element(elements=('Quản trị Nhân sự', 'Quản trị Kinh doanh', 'Tâm lý học'))
    elif 'Kinh doanh' in group or 'Sales' in group:
        major = fake.random_element(elements=('Kinh doanh', 'Quản trị Kinh doanh', 'Marketing'))
    else:
        major = fake.random_element(elements=('Quản trị Kinh doanh', 'Kinh tế', 'Kỹ thuật'))
    
    university = fake.random_element(elements=(
        'Đại học Bách khoa Hà Nội',
        'Đại học Kinh tế Quốc dân',
        'Đại học Quốc gia Hà Nội',
        'Đại học Bách khoa TP.HCM',
        'Đại học Kinh tế TP.HCM',
        'Đại học Công nghệ Thông tin'
    ))
    
    grad_year = 2024 - experience_years - fake.random_int(0, 2)
    doc.add_paragraph(f"{major} | {university} | Tốt nghiệp {grad_year}")
    doc.add_paragraph()  # Khoảng trống
    
    # ========== KỸ NĂNG ==========
    doc.add_paragraph().add_run("KỸ NĂNG").bold = True
    
    # Kỹ năng chuyên môn
    doc.add_paragraph("Kỹ năng chuyên môn:", style='List Bullet')
    for skill in skills[:5]:  # Lấy 5 kỹ năng đầu
        doc.add_paragraph(f"  • {skill}", style='List Bullet')
    
    # Kỹ năng mềm
    soft_skills = ['Giao tiếp tốt', 'Làm việc nhóm', 'Quản lý thời gian', 'Giải quyết vấn đề']
    doc.add_paragraph("Kỹ năng mềm:", style='List Bullet')
    for skill in soft_skills:
        doc.add_paragraph(f"  • {skill}", style='List Bullet')
    
    doc.add_paragraph()  # Khoảng trống
    
    # ========== CHỨNG CHỈ ==========
    doc.add_paragraph().add_run("CHỨNG CHỈ").bold = True
    
    # Tạo 1-2 chứng chỉ phù hợp
    num_certs = fake.random_int(1, 2)
    cert_names = []
    
    if 'Developer' in position or 'Engineer' in position:
        cert_names = ['AWS Certified Solutions Architect', 'Google Cloud Professional', 'Microsoft Azure']
    elif 'Marketing' in position:
        cert_names = ['Google Ads Certification', 'Facebook Blueprint', 'HubSpot Content Marketing']
    elif 'Kế toán' in position:
        cert_names = ['Chứng chỉ Kế toán viên', 'ACCA', 'CPA']
    elif 'HR' in position or 'Nhân sự' in position:
        cert_names = ['SHRM Certified Professional', 'HRCI PHR', 'Chứng chỉ Quản trị Nhân sự']
    else:
        cert_names = ['Chứng chỉ chuyên môn', 'Professional Certificate']
    
    for i in range(num_certs):
        cert = fake.random_element(elements=cert_names)
        year = fake.random_int(2020, 2024)
        doc.add_paragraph(f"• {cert} ({year})", style='List Bullet')
    
    doc.add_paragraph()  # Khoảng trống
    
    # ========== NGÔN NGỮ ==========
    doc.add_paragraph().add_run("NGÔN NGỮ").bold = True
    doc.add_paragraph("• Tiếng Việt: Bản ngữ", style='List Bullet')
    doc.add_paragraph("• Tiếng Anh: " + fake.random_element(elements=('Khá', 'Tốt', 'Rất tốt')), style='List Bullet')
    
    # Lưu file
    doc.save(filepath)
    print(f"✅ Đã tạo: {os.path.basename(filepath)}")


# Định nghĩa 10 lĩnh vực khác nhau
job_positions = [
    {
        'group': 'Công nghệ Thông tin',
        'profession': 'Lập trình viên',
        'position': 'Frontend Developer',
        'skills': ['React', 'Vue.js', 'TypeScript', 'HTML/CSS', 'JavaScript', 'Redux', 'Next.js']
    },
    {
        'group': 'Công nghệ Thông tin',
        'profession': 'Lập trình viên',
        'position': 'Backend Developer',
        'skills': ['Python', 'Django', 'Node.js', 'PostgreSQL', 'RESTful API', 'MongoDB', 'Redis']
    },
    {
        'group': 'Marketing/PR/Quảng cáo',
        'profession': 'Digital Marketing',
        'position': 'Digital Marketing Specialist',
        'skills': ['SEO/SEM', 'Google Ads', 'Facebook Ads', 'Content Marketing', 'Analytics', 'Social Media']
    },
    {
        'group': 'Kế toán/Tài chính',
        'profession': 'Kế toán',
        'position': 'Kế toán tổng hợp',
        'skills': ['Kế toán tổng hợp', 'Excel nâng cao', 'Phần mềm kế toán', 'Báo cáo tài chính', 'Khai báo thuế']
    },
    {
        'group': 'Nhân sự/Hành chính/Pháp chế',
        'profession': 'Nhân sự',
        'position': 'HR Manager',
        'skills': ['Tuyển dụng', 'Quản lý nhân sự', 'Đào tạo', 'HRIS', 'Chính sách nhân sự', 'Đánh giá hiệu suất']
    },
    {
        'group': 'Kinh doanh/Bán hàng',
        'profession': 'Sales IT/Phần mềm',
        'position': 'Sales IT/Phần mềm',
        'skills': ['Bán hàng B2B', 'Tư vấn giải pháp', 'CRM', 'Cloud Computing', 'Đàm phán', 'Quản lý khách hàng']
    },
    {
        'group': 'Công nghệ Thông tin',
        'profession': 'Data/AI',
        'position': 'Data Analyst',
        'skills': ['SQL', 'Python', 'Tableau', 'Excel', 'Power BI', 'Data Visualization', 'Statistics']
    },
    {
        'group': 'Marketing/PR/Quảng cáo',
        'profession': 'Brand Marketing',
        'position': 'Brand Manager',
        'skills': ['Brand Strategy', 'Market Research', 'Campaign Management', 'Brand Positioning', 'Marketing Mix']
    },
    {
        'group': 'Chăm sóc khách hàng/Vận hành',
        'profession': 'Operations',
        'position': 'Operations Manager',
        'skills': ['Operations Management', 'Process Optimization', 'Team Leadership', 'Supply Chain', 'Quality Control']
    },
    {
        'group': 'Kế toán/Tài chính',
        'profession': 'Tài chính',
        'position': 'Financial Analyst',
        'skills': ['Phân tích tài chính', 'Financial Modeling', 'Excel', 'Forecasting', 'Budgeting', 'Risk Analysis']
    }
]

if __name__ == '__main__':
    script_dir = os.path.dirname(os.path.abspath(__file__))
    cv_dir = os.path.join(script_dir, 'sample_cvs_docx')
    os.makedirs(cv_dir, exist_ok=True)
    
    print(f"📝 Tạo 10 CV mẫu dạng DOCX cho các lĩnh vực khác nhau...")
    print(f"📁 Lưu tại: {cv_dir}\n")
    
    for idx, job in enumerate(job_positions, 1):
        # Tạo thông tin cá nhân
        name = fake.name()
        email = f"{sanitize_filename(name.lower().replace(' ', '.'))}@example.com"
        phone = fake.phone_number()
        address = fake.address().replace('\n', ', ')
        experience_years = fake.random_int(2, 6)  # 2-6 năm kinh nghiệm
        
        # Tạo tên file
        filename = f"CV_{idx:02d}_{sanitize_filename(name)}_{sanitize_filename(job['position'])}.docx"
        filepath = os.path.join(cv_dir, filename)
        
        # Tạo CV
        create_cv_docx(
            filepath,
            name,
            email,
            phone,
            address,
            experience_years,
            job['skills'],
            job['position'],
            job['group'],
            job['profession']
        )
    
    print(f"\n✅ Đã tạo 10 CV mẫu dạng DOCX!")
    print(f"📁 Vị trí: {cv_dir}")
    print(f"\n📋 Danh sách CV đã tạo:")
    for idx, job in enumerate(job_positions, 1):
        print(f"  {idx}. {job['position']} ({job['group']})")

